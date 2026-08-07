from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document

PREVIOUS_HEADING_RE = re.compile(r"(?:^|\s)(?:1\.2\s*)?上一次(?:定期)?检测状况")
NEXT_MAIN_RE = re.compile(r"^\s*2(?:\.0)?\s*(?:检测目的|检查目的)")
PREVIOUS_BCI_RE = re.compile(r"(?:整体|总体)?技术状况指数\s*BCI\s*[=＝]?\s*(\d+(?:\.\d+)?)", re.I)
PREVIOUS_GRADE_RE = re.compile(r"整体技术状况等级(?:评定|评价|确定)?为\s*([A-Ea-e]\s*级|[一二三四五六]类)")
PREVIOUS_GRADE_FALLBACK_RE = re.compile(r"(?:总体|整体)技术状况[^。；;]{0,100}?(?:等级|级别)[^。；;]{0,30}?([A-Ea-e]\s*级|[一二三四五六]类)")
CAUSE_HEADING_RE = re.compile(r"病害(?:原因|成因)分析|病害原因")
SAFETY_RE = re.compile(r"安全性评估|安全评估|可以正常运营|正常运营|承载能力|安全隐患|耐久性")
HISTORY_HEADERS = {
    "location": ("位置", "部位", "结构部位"),
    "previous": ("上一次检测结果", "上次检测结果", "上一次定检结果", "历史检测结果"),
    "current": ("本次检测结果", "本次定检结果", "当前检测结果"),
    "development": ("发展状况", "发展情况", "变化情况", "病害发展"),
}
DISEASE_TERMS = (
    "渗水泛碱", "浸水泛碱", "破损露筋", "露筋锈蚀", "蜂窝麻面", "网状裂缝", "纵向裂缝", "横向裂缝",
    "贯穿裂缝", "裂缝", "破损", "剥落", "脱落", "露筋", "锈蚀", "渗水", "泛碱", "变形", "堵塞", "车辙",
    "坑槽", "磨损", "缺失", "沉降", "冲刷", "积水", "开裂", "腐蚀", "松动", "错位", "蜂窝", "麻面",
)
COMPONENTS = (
    ("上部结构", ("上部结构", "主梁", "梁", "板", "翼板", "腹板", "横隔", "铰缝", "索", "桥面板")),
    ("下部结构", ("下部结构", "桥墩", "墩", "桥台", "台帽", "台身", "盖梁", "支座", "基础", "挡块")),
    ("桥面系", ("桥面系", "桥面", "铺装", "路面", "伸缩缝", "护栏", "栏杆", "排水", "泄水", "人行道", "路缘")),
    ("主体结构", ("主体结构", "顶板", "侧墙", "墙体", "底板", "衬砌", "洞口", "翼墙")),
    ("附属设施", ("附属设施", "栏杆", "扶手", "照明", "标志", "排水")),
)


def clean(value: str) -> str:
    return " ".join((value or "").replace("\u3000", " ").split()).strip()


def compact(value: str) -> str:
    return re.sub(r"\s+", "", value or "")


def history_mapping(table: Any) -> tuple[int, dict[str, int]] | None:
    for r_idx, row in enumerate(table.rows[:4]):
        mapping: dict[str, int] = {}
        for c_idx, cell in enumerate(row.cells):
            value = compact(cell.text)
            for field, aliases in HISTORY_HEADERS.items():
                if field not in mapping and any(compact(alias) in value for alias in aliases):
                    mapping[field] = c_idx
                    break
        if {"previous", "current", "development"}.issubset(mapping):
            return r_idx, mapping
    return None


def history_groups(doc: Document) -> list[dict[str, str]]:
    groups: list[dict[str, str]] = []
    for table in doc.tables:
        hm = history_mapping(table)
        if hm is None:
            continue
        hrow, mapping = hm
        for row in table.rows[hrow + 1 :]:
            def cell(field: str) -> str:
                idx = mapping.get(field)
                return clean(row.cells[idx].text) if idx is not None and idx < len(row.cells) else ""
            item = {k: cell(k) for k in ("location", "previous", "current", "development")}
            if any(item.values()):
                groups.append(item)
    return groups


def previous_facts(doc: Document) -> tuple[str, str, str]:
    paras = [clean(p.text) for p in doc.paragraphs]
    start = next((i for i, t in enumerate(paras) if PREVIOUS_HEADING_RE.search(t)), None)
    if start is None:
        return "", "", ""
    section: list[str] = []
    for text in paras[start + 1 :]:
        if NEXT_MAIN_RE.search(text):
            break
        if text:
            section.append(text)
        if len(section) >= 30:
            break
    joined = " ".join(section)
    score = PREVIOUS_BCI_RE.search(joined)
    grade = PREVIOUS_GRADE_RE.search(joined) or PREVIOUS_GRADE_FALLBACK_RE.search(joined)
    no_previous = "无上一次检测" if any("无上一次检测" in t or "无上次检测" in t for t in section) else ""
    return score.group(1) if score else "", grade.group(1).replace(" ", "") if grade else "", no_previous


def cause_paragraphs(doc: Document) -> list[str]:
    paras = [clean(p.text) for p in doc.paragraphs]
    starts = [i for i, t in enumerate(paras) if CAUSE_HEADING_RE.search(t)]
    if not starts:
        return []
    start = starts[0]
    result: list[str] = []
    # 7.1.2 is commonly followed by disease-specific subheadings.  They belong
    # to the same cause-analysis section; stop only at the next major section.
    for t in paras[start + 1 : start + 60]:
        if re.match(r"^\s*(?:7\.[2-9](?:\.|\s)|8(?:\.|\s))", t):
            break
        if any(marker in t for marker in ("由于", "导致", "原因", "所致", "引起", "主要是", "主要为", "与", "受")):
            result.append(t)
    return result

def safety_paragraphs(doc: Document) -> list[str]:
    return [clean(p.text) for p in doc.paragraphs if p.text and SAFETY_RE.search(clean(p.text))][:20]


def record_component(defect: dict[str, Any]) -> str:
    text = clean(" ".join(str(defect.get(k, "")) for k in ("location", "defect_type", "description")))
    for key, terms in COMPONENTS:
        if any(term in text for term in terms):
            return key
    return ""


def disease_terms(text: str) -> list[str]:
    found = [term for term in DISEASE_TERMS if term in text]
    found.sort(key=len, reverse=True)
    kept: list[str] = []
    for term in found:
        if not any(term in x for x in kept):
            kept.append(term)
    return kept[:6]


def history_location_tokens(value: str) -> tuple[str, ...]:
    text = clean(value)
    tokens = re.findall(
        r"(?:左幅|右幅)?\d+#(?:跨|墩|台|缝|支座|梁|板)|"
        r"(?:距)?\d+#(?:伸缩缝|墩|台)|K\d+(?:\+\d+)?",
        text,
        flags=re.IGNORECASE,
    )
    if tokens:
        return tuple(dict.fromkeys(tokens))
    compact_text = compact(text)
    generic = {"桥面", "主梁", "桥墩", "桥台", "支座", "护栏", "栏杆", "伸缩缝", "路面", "上部结构", "下部结构", "桥面系"}
    if 3 <= len(compact_text) <= 24 and compact_text not in generic:
        return (compact_text,)
    return ()


def history_location_match(defect: dict[str, Any], text: str) -> bool:
    tokens = history_location_tokens(str(defect.get("location", "")))
    if not tokens:
        return False
    target = compact(text)
    return any(compact(token) in target for token in tokens)


def best_group(defect: dict[str, Any], groups: list[dict[str, str]]) -> dict[str, str] | None:
    component = record_component(defect)
    terms = disease_terms(clean(str(defect.get("defect_type", "")) + " " + str(defect.get("description", ""))))
    ranked: list[tuple[int, int, dict[str, str]]] = []
    for order, group in enumerate(groups):
        group_text = " ".join(group.get(k, "") for k in ("current", "development", "previous"))
        disease_hits = sum(1 for term in terms if term in group_text)
        dtype = str(defect.get("defect_type", ""))
        if not disease_hits and not (dtype and dtype in group_text):
            continue
        score = 0
        if component and component in group.get("location", ""):
            score += 8
        if history_location_match(defect, group_text):
            score += 12
        score += 3 * disease_hits
        if dtype and dtype in group_text:
            score += 5
        if score:
            ranked.append((score, -order, group))
    return max(ranked, default=(0, 0, None), key=lambda x: (x[0], x[1]))[2]


def would_enrich(defect: dict[str, Any], groups: list[dict[str, str]]) -> bool:
    if str(defect.get("previous_status", "")) not in ("", "无") or str(defect.get("development", "")) not in ("", "无") or str(defect.get("is_new", "")) not in ("", "否"):
        return False
    group = best_group(defect, groups)
    if not group:
        return False
    terms = disease_terms(str(defect.get("defect_type", "")) + " " + str(defect.get("description", "")))
    previous = group.get("previous", "")
    current = group.get("current", "")
    development = group.get("development", "")
    previous_match = bool(terms) and history_location_match(defect, previous) and any(term in previous for term in terms)
    current_match = bool(terms) and history_location_match(defect, current) and any(term in current for term in terms)
    development_match = bool(terms) and any(term in development for term in terms)
    if previous_match:
        return True
    if "新增" in development and current_match and development_match:
        return True
    if development_match and current_match and any(term in development for term in ("发展", "加重", "扩大", "修复", "减轻", "变化")):
        return True
    return False


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("evidence_root", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args()
    root = args.evidence_root
    manifest = json.loads((root / "metadata" / "sample_manifest.json").read_text(encoding="utf-8"))
    report32 = json.loads((root / "logs" / "32-80__prediction-report.json").read_text(encoding="utf-8"))

    retrieval = Counter()
    retrieval_examples: list[dict[str, str]] = []
    for rec in report32.get("records", []):
        sid = str(rec.get("sample_id", ""))
        narrative = ((rec.get("semantic_trace") or {}).get("narrative") or {})
        for task, task_rec in (narrative.get("retrieval_by_task") or {}).items():
            for hit in (task_rec.get("hits") or []):
                if hit.get("source_bucket") != "report_evidence":
                    continue
                retrieval["report_hits"] += 1
                if str(hit.get("sample_id", "")) == sid:
                    retrieval["same_sample"] += 1
                else:
                    retrieval["cross_sample"] += 1
                    if len(retrieval_examples) < 8:
                        retrieval_examples.append({
                            "sample_id": sid,
                            "task": str(task),
                            "hit_sample_id": str(hit.get("sample_id", "")),
                            "text": clean(str(hit.get("text", "")))[:180],
                        })

    rows: list[dict[str, Any]] = []
    totals = Counter()
    for item in manifest:
        candidates = sorted((root / "samples" / "source-docs").glob(f"{int(item['index']):02d}__*.docx"))
        if len(candidates) != 1:
            raise RuntimeError(f"cannot resolve source DOCX for sample index {item['index']}: {candidates}")
        source = candidates[0]
        doc = Document(source)
        score, grade, no_previous = previous_facts(doc)
        groups = history_groups(doc)
        causes = cause_paragraphs(doc)
        safety = safety_paragraphs(doc)
        pred32 = item["runs"]["32-80"]["prediction"]
        defects = pred32.get("defects", [])
        enrich_count = sum(1 for defect in defects if would_enrich(defect, groups))
        old_history_nondefault = sum(
            1 for d in defects
            if str(d.get("previous_status", "")) not in ("", "无")
            or str(d.get("development", "")) not in ("", "无")
            or str(d.get("is_new", "")) not in ("", "否")
        )
        row = {
            "sample_id": item["sample_id"],
            "previous_score_source": score,
            "previous_grade_source": grade,
            "explicit_no_previous": bool(no_previous),
            "old_32_previous_score": pred32.get("summary", {}).get("previous_overall_score", ""),
            "old_32_previous_grade": pred32.get("summary", {}).get("previous_overall_grade", ""),
            "history_group_count": len(groups),
            "cause_evidence_paragraphs": len(causes),
            "safety_evidence_paragraphs": len(safety),
            "defect_count": len(defects),
            "old_history_nondefault_rows": old_history_nondefault,
            "conservative_history_rows_patch_can_enrich": enrich_count,
        }
        rows.append(row)
        if score: totals["previous_score_found"] += 1
        if grade: totals["previous_grade_found"] += 1
        if no_previous: totals["explicit_no_previous"] += 1
        if groups: totals["history_table_samples"] += 1
        if causes: totals["cause_section_samples"] += 1
        if safety: totals["safety_evidence_samples"] += 1
        totals["defects"] += len(defects)
        totals["history_rows_patch_can_enrich"] += enrich_count
        totals["old_history_nondefault_rows"] += old_history_nondefault

    output = {
        "sample_count": len(rows),
        "totals": dict(totals),
        "retrieval_32_80": dict(retrieval),
        "retrieval_cross_sample_rate": (retrieval["cross_sample"] / retrieval["report_hits"] if retrieval["report_hits"] else 0),
        "retrieval_examples": retrieval_examples,
        "samples": rows,
    }
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"totals": output["totals"], "retrieval_32_80": output["retrieval_32_80"], "cross_rate": output["retrieval_cross_sample_rate"]}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
