"""Build the small public derivative dataset from locally owned label files.

The script reads only label .docx files and writes structured JSONL. It never
copies source reports into the output directory.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document


TRAIN_NAMES = [
    "柏杨湾大桥",
    "桂花新村大桥",
    "茶亭大桥",
    "大堡桥",
    "大田坝大桥",
    "道角村大桥",
    "丁家院大桥",
    "果园大桥",
    "黄桷湾主线桥",
    "鸡冠嘴中桥",
    "界石立交主线VI号桥",
    "成渝2号人行天桥",
    "上界路K40+924人行天桥检测报告",
    "华岩寺大桥",
    "南山立交A-1匝道桥",
    "四公里连接线桥",
]

VALIDATION_NAMES = [
    "12-025杨公桥立交DA1匝道桥",
    "12-027杨公桥立交DA-ED匝道桥",
    "12-028杨公桥DB-DC匝道桥",
    "12-030杨公桥DC匝道桥",
    "12-031杨公桥AC匝道桥",
    "12-035杨公桥立交EC匝道桥",
]


def clean(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\u00a0", " ")).strip()


def norm(value: str) -> str:
    value = value.lower().replace("无对比年度的信息提取报告", "")
    value = value.replace("无对比年度定检报告", "")
    value = value.replace("报告", "")
    return re.sub(r"[^0-9a-z\u4e00-\u9fff]+", "", value)


def label_base(path: Path) -> str:
    return re.sub(
        r"-?无对比年度(?:的信息提取报告|定检报告)$", "", path.stem
    ).strip()


def locate_label(labels_dir: Path, wanted: str) -> Path:
    matches = [p for p in labels_dir.rglob("*.docx") if norm(label_base(p)) == norm(wanted)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one label for {wanted!r}, found {matches}")
    return matches[0]


def locate_report(reports_dir: Path, label_name: str) -> Path | None:
    target = norm(label_name)
    exact = []
    partial = []
    for path in reports_dir.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in {".doc", ".docx"}:
            continue
        current = norm(path.stem)
        if current == target:
            exact.append(path)
        elif target in current or current in target:
            partial.append(path)
    candidates = exact or partial
    return sorted(candidates, key=lambda p: (len(p.name), str(p)))[0] if candidates else None


def paragraph_blocks(document: Document) -> list[str]:
    return [clean(p.text) for p in document.paragraphs if clean(p.text)]


def between(paragraphs: list[str], start: str, end: str | None = None) -> list[str]:
    try:
        start_index = paragraphs.index(start) + 1
    except ValueError:
        return []
    end_index = len(paragraphs)
    if end is not None:
        try:
            end_index = paragraphs.index(end, start_index)
        except ValueError:
            pass
    return paragraphs[start_index:end_index]


def extract(label_path: Path, labels_dir: Path, reports_dir: Path, split: str) -> dict:
    document = Document(str(label_path))
    paragraphs = paragraph_blocks(document)
    if len(document.tables) < 3:
        raise RuntimeError(f"Expected three tables in {label_path}")

    summary: dict[str, str] = {}
    for row in document.tables[0].rows[1:]:
        cells = [clean(cell.text) for cell in row.cells]
        if len(cells) >= 2 and cells[0]:
            summary[cells[0]] = cells[1]

    recommendations = []
    for row in document.tables[1].rows[1:]:
        cells = [clean(cell.text) for cell in row.cells]
        if any(cells):
            cells += [""] * (4 - len(cells))
            recommendations.append(
                {
                    "index": cells[0],
                    "category": cells[1],
                    "content": cells[2],
                    "location": cells[3],
                }
            )

    defects = []
    for row in document.tables[2].rows[1:]:
        cells = [clean(cell.text) for cell in row.cells]
        if any(cells):
            cells += [""] * (7 - len(cells))
            defects.append(
                {
                    "index": cells[0],
                    "location": cells[1],
                    "defect_type": cells[2],
                    "description": cells[3],
                    "is_new": cells[4],
                    "previous_status": cells[5],
                    "development": cells[6],
                }
            )

    name = label_base(label_path)
    source_report = locate_report(reports_dir, name)
    sample_id = f"{label_path.parent.name}-{name}"
    unique_types = sorted({d["defect_type"] for d in defects if d["defect_type"]})
    return {
        "sample_id": sample_id,
        "split": split,
        "summary": {
            "bridge_name": summary.get("桥梁名称", ""),
            "report_date": summary.get("报告日期", ""),
            "overall_score": summary.get("总体评分", ""),
            "overall_grade": summary.get("总体等级", ""),
            "superstructure_score": summary.get("上部结构评分", ""),
            "superstructure_grade": summary.get("上部结构等级", ""),
            "substructure_score": summary.get("下部结构评分", ""),
            "substructure_grade": summary.get("下部结构等级", ""),
            "deck_score": summary.get("桥面系评分", ""),
            "deck_grade": summary.get("桥面系等级", ""),
            "previous_overall_score": summary.get("上一次总体评分", ""),
            "previous_overall_grade": summary.get("上一次总体等级", ""),
            "trend": summary.get("病害发展趋势与具体说明", ""),
            "overall_conclusion": summary.get("总体结论", ""),
            "risk_points": summary.get("主要风险点", ""),
            "recommendations_summary": summary.get("建议", ""),
        },
        "detailed_conclusion": between(paragraphs, "（1）详细结论", "（2）建议明细"),
        "recommendations": recommendations,
        "defects": defects,
        "causes": between(paragraphs, "病害成因", "处置建议"),
        "treatments": between(paragraphs, "处置建议", "安全影响"),
        "safety_impact": between(paragraphs, "安全影响"),
        "statistics": {
            "label_character_count": sum(len(p) for p in paragraphs),
            "recommendation_count": len(recommendations),
            "defect_count": len(defects),
            "defect_type_count": len(unique_types),
            "defect_types": unique_types,
        },
        "provenance": {
            "label_relative_path": str(label_path.relative_to(labels_dir)).replace("\\", "/"),
            "source_report_relative_path": (
                str(source_report.relative_to(reports_dir)).replace("\\", "/")
                if source_report else None
            ),
            "raw_report_included": False,
            "derivation": "Structured from a locally held competition training-label document; raw source files are not redistributed.",
        },
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--labels-dir", type=Path, required=True)
    parser.add_argument("--reports-dir", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()

    selected = [("train", name) for name in TRAIN_NAMES] + [
        ("validation", name) for name in VALIDATION_NAMES
    ]
    args.output_dir.mkdir(parents=True, exist_ok=True)
    records = []
    for split, name in selected:
        label_path = locate_label(args.labels_dir, name)
        records.append(extract(label_path, args.labels_dir, args.reports_dir, split))

    records.sort(key=lambda record: (record["split"], record["sample_id"]))
    for split in ("train", "validation"):
        output = args.output_dir / f"{split}.jsonl"
        with output.open("w", encoding="utf-8", newline="\n") as handle:
            for record in records:
                if record["split"] == split:
                    handle.write(json.dumps(record, ensure_ascii=False, separators=(",", ":")) + "\n")

    manifest = {
        "dataset_name": "city-infrastructure-inspection-rag-core",
        "dataset_version": "core-v1",
        "selection": {
            "train_count": sum(r["split"] == "train" for r in records),
            "validation_count": sum(r["split"] == "validation" for r in records),
            "train_source_group": "2012年",
            "validation_source_group": "2013年",
            "note": "The validation group is concentrated in Yangongqiao interchange facilities and is not a strict facility-disjoint generalization benchmark.",
        },
        "records": [
            {
                "sample_id": r["sample_id"],
                "split": r["split"],
                "bridge_name": r["summary"]["bridge_name"],
                "report_date": r["summary"]["report_date"],
                "defect_count": r["statistics"]["defect_count"],
                "source_report_relative_path": r["provenance"]["source_report_relative_path"],
            }
            for r in records
        ],
    }
    (args.output_dir / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    print(json.dumps({"train": manifest["selection"]["train_count"], "validation": manifest["selection"]["validation_count"]}, ensure_ascii=False))


if __name__ == "__main__":
    main()
