"""Minimal parser for the three-table Word label format."""

from __future__ import annotations

from pathlib import Path
import re
from typing import Iterable

from docx import Document

from ..audit.core import label_base, label_quality_flags, relative_path


class LabelParseError(RuntimeError):
    """A label was present but could not be converted into Gold JSON."""

    def __init__(self, code: str, message: str) -> None:
        self.code = code
        super().__init__(message)


SUMMARY_FIELDS = {
    "桥梁名称": "bridge_name",
    "桥梁编号": "bridge_id",
    "报告日期": "report_date",
    "检测日期": "report_date",
    "总体评分": "overall_score",
    "总体等级": "overall_grade",
    "上部结构评分": "superstructure_score",
    "上部结构等级": "superstructure_grade",
    "下部结构评分": "substructure_score",
    "下部结构等级": "substructure_grade",
    "桥面系评分": "deck_score",
    "桥面系等级": "deck_grade",
    "上一次总体评分": "previous_overall_score",
    "上一次总体等级": "previous_overall_grade",
    "病害发展趋势与具体说明": "trend",
    "总体结论": "overall_conclusion",
    "主要风险点": "risk_points",
    "建议": "recommendations_summary",
}


def clean(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\u00a0", " ")).strip()


def _row_cells(row: object) -> list[str]:
    return [clean(cell.text) for cell in row.cells]  # type: ignore[attr-defined]


def _looks_like_header(cells: list[str], markers: Iterable[str]) -> bool:
    joined = "|".join(cells).replace(" ", "")
    return all(marker.replace(" ", "") in joined for marker in markers)


def _header_index(table: object, marker_groups: Iterable[Iterable[str]]) -> int | None:
    rows = [_row_cells(row) for row in table.rows]  # type: ignore[attr-defined]
    for index, cells in enumerate(rows[:8]):
        if any(_looks_like_header(cells, markers) for markers in marker_groups):
            return index
    return None


def _find_table(document: object, marker_groups: Iterable[Iterable[str]], label: str) -> object:
    matches = []
    for table in document.tables:  # type: ignore[attr-defined]
        header_index = _header_index(table, marker_groups)
        if header_index is not None:
            matches.append((table, header_index))
    if not matches:
        raise LabelParseError("missing_%s_table" % label, f"cannot locate {label} table by headers")
    if len(matches) > 1:
        # Prefer the table with the most rows. This is deterministic and avoids
        # selecting a short explanatory table with similar words.
        matches.sort(key=lambda item: len(item[0].rows), reverse=True)  # type: ignore[attr-defined]
    return matches[0][0]


def _data_rows(table: object, marker_groups: Iterable[Iterable[str]]) -> list[list[str]]:
    rows = [_row_cells(row) for row in table.rows]  # type: ignore[attr-defined]
    header_index = _header_index(table, marker_groups)
    return rows[(header_index + 1) if header_index is not None else 0 :]


def _summary_rows(document: object) -> dict[str, str]:
    table = _find_table(
        document,
        (("字段", "内容"), ("桥梁名称",), ("总体评分",)),
        "summary",
    )
    rows = _data_rows(table, (("字段", "内容"), ("项目", "值")))
    values: dict[str, str] = {}
    for cells in rows:
        # Official labels are usually 3 columns (field/value/note), but some
        # variants use two field/value pairs in four columns.
        if len(cells) >= 3 and _normalise_summary_key(cells[0]) in SUMMARY_FIELDS:
            values[cells[0]] = cells[1]
            continue
        for index in range(0, len(cells) - 1, 2):
            key = cells[index].rstrip("：:").strip()
            value = cells[index + 1]
            if key:
                values[key] = value
    return values


def _normalise_summary_key(key: str) -> str:
    return clean(key).rstrip("：:").replace(" ", "")


def _summary(document: object) -> dict[str, str]:
    raw = _summary_rows(document)
    result = {
        field: ""
        for field in (
            "bridge_name",
            "bridge_id",
            "report_date",
            "overall_score",
            "overall_grade",
            "superstructure_score",
            "superstructure_grade",
            "substructure_score",
            "substructure_grade",
            "deck_score",
            "deck_grade",
            "previous_overall_score",
            "previous_overall_grade",
            "trend",
            "overall_conclusion",
            "risk_points",
            "recommendations_summary",
        )
    }
    for key, value in raw.items():
        field = SUMMARY_FIELDS.get(_normalise_summary_key(key))
        if field:
            result[field] = value
    return result


def _recommendations(document: object) -> list[dict[str, str]]:
    table = _find_table(
        document,
        (("建议类别", "建议内容"), ("维修建议", "病害部位")),
        "recommendations",
    )
    rows = _data_rows(table, (("建议类别", "建议内容"), ("维修建议", "病害部位")))
    result: list[dict[str, str]] = []
    for cells in rows:
        cells = (cells + [""] * 4)[:4]
        if any(cells):
            result.append(
                {
                    "index": cells[0],
                    "category": cells[1],
                    "content": cells[2],
                    "location": cells[3],
                }
            )
    return result


def _defects(document: object) -> list[dict[str, str]]:
    table = _find_table(
        document,
        (("病害部位", "病害类型", "病害描述"),),
        "defects",
    )
    rows = _data_rows(table, (("病害部位", "病害类型", "病害描述"),))
    result: list[dict[str, str]] = []
    for cells in rows:
        cells = (cells + [""] * 7)[:7]
        if any(cells):
            result.append(
                {
                    "index": cells[0],
                    "location": cells[1],
                    "defect_type": cells[2],
                    "description": cells[3],
                    "is_new": cells[4],
                    "previous_status": cells[5],
                    "development": cells[6],
                }
            )
    return result


def _paragraphs(document: object) -> list[str]:
    return [clean(paragraph.text) for paragraph in document.paragraphs if clean(paragraph.text)]  # type: ignore[attr-defined]


def _heading_index(paragraphs: list[str], markers: Iterable[str], start: int = 0) -> int | None:
    for index in range(start, len(paragraphs)):
        text = paragraphs[index].replace(" ", "")
        if any(marker in text for marker in markers):
            return index
    return None


def _section(paragraphs: list[str], starts: Iterable[str], ends: Iterable[str] = ()) -> list[str]:
    start = _heading_index(paragraphs, starts)
    if start is None:
        return []
    end = len(paragraphs)
    for index in range(start + 1, len(paragraphs)):
        text = paragraphs[index].replace(" ", "")
        if any(marker in text for marker in ends):
            end = index
            break
    return paragraphs[start + 1 : end]


def _split_from_path(label_path: Path, labels_root: Path) -> str:
    parts = [part.casefold() for part in label_path.relative_to(labels_root).parts]
    if any("2013" in part for part in parts):
        return "validation"
    if any("2012" in part for part in parts):
        return "train"
    return "unknown"


def parse_label_docx(
    label_path: Path | str,
    labels_root: Path | str,
    source_report_relative_path: str | None = None,
) -> dict[str, object]:
    """Parse one DOCX label and return only relative-path provenance."""

    label_path = Path(label_path)
    labels_root = Path(labels_root)
    if label_path.suffix.casefold() == ".doc":
        raise LabelParseError(
            "legacy_doc_unsupported",
            "legacy .doc labels must be converted to .docx before parsing",
        )
    if label_path.suffix.casefold() != ".docx":
        raise LabelParseError("unsupported_extension", f"unsupported label extension: {label_path.suffix}")
    try:
        document = Document(str(label_path))
    except Exception as exc:
        raise LabelParseError("docx_read_failed", f"cannot read label {label_path.name}: {exc}") from exc
    paragraphs = _paragraphs(document)
    recommendations = _recommendations(document)
    defects = _defects(document)
    unique_types = sorted({item["defect_type"] for item in defects if item["defect_type"]})
    relative_label = relative_path(label_path, labels_root)
    label_name = label_base(label_path.stem)
    parent = label_path.relative_to(labels_root).parent.as_posix()
    sample_id = label_name if parent == "." else f"{parent.replace('/', '-')}-{label_name}"
    record = {
        "sample_id": sample_id,
        "split": _split_from_path(label_path, labels_root),
        "summary": _summary(document),
        "detailed_conclusion": _section(paragraphs, ("详细结论",), ("建议明细",)),
        "recommendations": recommendations,
        "defects": defects,
        "causes": _section(paragraphs, ("病害成因",), ("处置建议",)),
        "treatments": _section(paragraphs, ("处置建议",), ("安全影响",)),
        "safety_impact": _section(paragraphs, ("安全影响",)),
        "statistics": {
            "label_character_count": sum(len(paragraph) for paragraph in paragraphs),
            "recommendation_count": len(recommendations),
            "defect_count": len(defects),
            "defect_type_count": len(unique_types),
            "defect_types": unique_types,
        },
        "provenance": {
            "label_relative_path": relative_label,
            "source_report_relative_path": source_report_relative_path,
            "raw_report_included": False,
            "derivation": "Structured from a locally held label DOCX; raw source files are not redistributed.",
        },
    }
    record["quality_flags"] = label_quality_flags(record)
    return record
