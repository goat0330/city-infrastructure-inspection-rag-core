"""Small, JSON-safe validator for rendered inspection-report DOCX files."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document


_ROLES = ("summary", "recommendations", "defects")

_HEADING_ALIASES = {
    "summary": ("桥梁概要", "桥梁概况", "概要表", "简要信息", "基本信息", "详细结论", "概要"),
    "recommendations": ("建议明细", "建议列表", "维修建议", "建议"),
    "defects": ("病害列表", "病害明细", "病害清单", "病害"),
}

_TABLE_MARKERS = {
    "summary": (("字段", "内容"), ("项目", "值"), ("概要", "内容"), ("桥梁名称",), ("总体评分",)),
    "recommendations": (
        ("建议类别", "建议内容"),
        ("维修建议", "病害部位"),
        ("建议内容", "病害部位"),
        ("建议", "内容"),
    ),
    "defects": (("病害部位", "病害类型", "病害描述"),),
}

_FIELD_ALIASES = {
    "recommendations": (
        ("index", ("序号", "编号")),
        ("category", ("建议类别", "类别", "维修建议")),
        ("content", ("建议内容", "内容")),
        ("location", ("病害部位", "部位", "位置")),
    ),
    "defects": (
        ("index", ("序号", "编号")),
        ("location", ("病害部位", "部位", "位置")),
        ("defect_type", ("病害类型", "类型")),
        ("description", ("病害描述", "描述")),
        ("is_new", ("是否新增", "新增")),
        ("previous_status", ("历史状态", "上一次定检状态", "既有病害状态")),
        ("development", ("发展", "发展程度", "发展趋势")),
    ),
}

_SUMMARY_KEYS = {
    "桥梁名称",
    "桥梁编号",
    "报告日期",
    "检测日期",
    "总体评分",
    "总体等级",
    "上部结构评分",
    "上部结构等级",
    "下部结构评分",
    "下部结构等级",
    "桥面系评分",
    "桥面系等级",
    "上一次总体评分",
    "上一次总体等级",
    "病害发展趋势与具体说明",
    "总体结论",
    "主要风险点",
    "建议",
}

_HEADING_PREFIX = re.compile(r"^[\s\u3000\(\)（）\[\]【】\d０-９IVXivx]+")


def validate_submission(path: str | Path) -> dict[str, Any]:
    """Validate one DOCX and return a deterministic JSON-serializable report.

    The report never contains the caller's path.  It records the file name,
    fixed-order checks, explicit failures, and row/empty-field metrics only.
    """

    source_path = Path(path)
    file_name = source_path.name
    checks: list[dict[str, Any]] = []
    failures: list[dict[str, Any]] = []
    metrics: dict[str, Any] = {
        "table_count": 0,
        "summary_table_index": None,
        "recommendations_table_index": None,
        "defects_table_index": None,
        "recommendation_rows": 0,
        "defect_rows": 0,
        "row_counts": {"recommendations": 0, "defects": 0},
        "empty_field_count": 0,
        "empty_fields": [],
    }

    extension_ok = source_path.suffix.casefold() == ".docx"
    temporary = _is_temporary_name(file_name)
    filename_ok = extension_ok and not temporary
    _add_check(
        checks,
        "filename",
        filename_ok,
        "DOCX extension and non-temporary filename are required.",
        {
            "extension": source_path.suffix.casefold(),
            "is_docx": extension_ok,
            "is_temporary": temporary,
        },
    )
    if not extension_ok:
        _add_failure(failures, "invalid_extension", "The submission filename must end with .docx.")
    if temporary:
        _add_failure(failures, "temporary_file", "Temporary Word files cannot be submitted.")

    document = None
    try:
        if not source_path.is_file():
            raise FileNotFoundError
        document = Document(str(source_path))
    except Exception:
        _add_check(
            checks,
            "docx_openable",
            False,
            "The file cannot be opened by python-docx.",
            {"opened": False},
        )
        _add_failure(failures, "docx_read_failed", "The file cannot be opened by python-docx.")
    else:
        _add_check(
            checks,
            "docx_openable",
            True,
            "The file opened successfully with python-docx.",
            {"opened": True},
        )

    if document is None:
        _add_not_checked(checks, "headings", "Document headings could not be inspected.")
        for role in _ROLES:
            _add_not_checked(checks, f"{role}_table", "Document tables could not be inspected.")
        _add_not_checked(checks, "recommendation_rows", "Recommendation rows could not be counted.")
        _add_not_checked(checks, "defect_rows", "Defect rows could not be counted.")
        _add_not_checked(checks, "empty_fields", "Table fields could not be inspected.")
    else:
        headings = _heading_results(document)
        missing_headings = [role for role in _ROLES if not headings[role]["found"]]
        _add_check(
            checks,
            "headings",
            not missing_headings,
            "Required section headings are present." if not missing_headings else "Required section headings are missing.",
            {"required": list(_ROLES), "found": headings},
        )
        for role in missing_headings:
            _add_failure(
                failures,
                "missing_heading",
                f"The required {role} heading is missing.",
                {"section": role},
            )

        metrics["table_count"] = len(document.tables)
        tables = _find_tables(document)
        empty_fields: list[dict[str, Any]] = []
        row_counts: dict[str, int] = {}
        for role in _ROLES:
            table_info = tables.get(role)
            if table_info is None:
                metrics[f"{role}_table_index"] = None
                _add_check(
                    checks,
                    f"{role}_table",
                    False,
                    f"The {role} table structure is missing.",
                    {"table_index": None},
                )
                _add_failure(
                    failures,
                    f"missing_{role}_table",
                    f"The {role} table structure is missing.",
                )
                row_counts[role] = 0
                continue

            table_index = table_info["table_index"]
            metrics[f"{role}_table_index"] = table_index
            _add_check(
                checks,
                f"{role}_table",
                True,
                f"The {role} table structure is present.",
                {
                    "table_index": table_index,
                    "header_row": table_info["header_row"],
                },
            )
            row_count, table_empty_fields = _inspect_table_rows(
                role,
                table_info,
                empty_fields,
                failures,
            )
            row_counts[role] = row_count

        metrics["recommendation_rows"] = row_counts.get("recommendations", 0)
        metrics["defect_rows"] = row_counts.get("defects", 0)
        metrics["row_counts"] = {
            "recommendations": metrics["recommendation_rows"],
            "defects": metrics["defect_rows"],
        }
        _add_check(
            checks,
            "recommendation_rows",
            metrics["recommendation_rows"] > 0,
            "At least one recommendation data row is required."
            if metrics["recommendation_rows"] == 0
            else "Recommendation data rows were counted.",
            {"count": metrics["recommendation_rows"], "minimum": 1},
        )
        if metrics["recommendation_rows"] == 0 and "recommendations" in tables:
            _add_failure(
                failures,
                "no_recommendation_rows",
                "The recommendations table has no data rows.",
            )
        _add_check(
            checks,
            "defect_rows",
            metrics["defect_rows"] > 0,
            "At least one defect data row is required."
            if metrics["defect_rows"] == 0
            else "Defect data rows were counted.",
            {"count": metrics["defect_rows"], "minimum": 1},
        )
        if metrics["defect_rows"] == 0 and "defects" in tables:
            _add_failure(
                failures,
                "no_defect_rows",
                "The defects table has no data rows.",
            )

        empty_fields.sort(key=lambda item: (item["section"], item["row"], item["column"]))
        metrics["empty_fields"] = empty_fields
        metrics["empty_field_count"] = len(empty_fields)
        _add_check(
            checks,
            "empty_fields",
            not empty_fields,
            "Required table fields are populated."
            if not empty_fields
            else "One or more required table fields are empty.",
            {"count": len(empty_fields)},
        )

    valid = not failures and all(check["passed"] for check in checks)
    return {
        "status": "passed" if valid else "failed",
        "valid": valid,
        "file_name": file_name,
        "checks": checks,
        "failures": failures,
        "metrics": metrics,
    }


def _add_check(
    checks: list[dict[str, Any]],
    name: str,
    passed: bool,
    message: str,
    details: dict[str, Any],
) -> None:
    checks.append(
        {
            "name": name,
            "status": "passed" if passed else "failed",
            "passed": passed,
            "message": message,
            "details": details,
        }
    )


def _add_not_checked(checks: list[dict[str, Any]], name: str, message: str) -> None:
    checks.append(
        {
            "name": name,
            "status": "not_checked",
            "passed": False,
            "message": message,
            "details": {},
        }
    )


def _add_failure(
    failures: list[dict[str, Any]],
    code: str,
    message: str,
    details: dict[str, Any] | None = None,
) -> None:
    failure: dict[str, Any] = {"code": code, "message": message}
    if details is not None:
        failure["details"] = details
    failures.append(failure)


def _is_temporary_name(file_name: str) -> bool:
    lowered = file_name.casefold()
    return lowered.startswith("~$") or lowered.startswith(".~lock.") or lowered.startswith(".~")


def _heading_results(document: Any) -> dict[str, dict[str, Any]]:
    paragraphs = [str(paragraph.text or "") for paragraph in document.paragraphs]
    results: dict[str, dict[str, Any]] = {}
    for role in _ROLES:
        matched = None
        for paragraph in paragraphs:
            for alias in _HEADING_ALIASES[role]:
                if _heading_matches(paragraph, alias):
                    matched = paragraph.strip()
                    break
            if matched is not None:
                break
        results[role] = {"found": matched is not None, "matched": matched}
    return results


def _heading_matches(text: str, alias: str) -> bool:
    normalized = _normalise(text)
    normalized = _HEADING_PREFIX.sub("", normalized)
    marker = _normalise(alias)
    if marker in {"建议", "病害", "概要"}:
        return normalized == marker
    return marker in normalized


def _find_tables(document: Any) -> dict[str, dict[str, Any]]:
    candidates: dict[str, list[tuple[int, int]]] = {role: [] for role in _ROLES}
    for table_index, table in enumerate(document.tables):
        for role in _ROLES:
            for header_row, cells in _header_rows(table, role):
                if any(_contains_markers(cells, group) for group in _TABLE_MARKERS[role]):
                    candidates[role].append((table_index, header_row))
                    break

    selected: dict[str, dict[str, Any]] = {}
    used_tables: set[int] = set()
    for role in _ROLES:
        ordered = sorted(
            candidates[role],
            key=lambda item: (-len(document.tables[item[0]].rows), item[0], item[1]),
        )
        for table_index, header_row in ordered:
            if table_index in used_tables:
                continue
            selected[role] = {
                "table_index": table_index,
                "header_row": header_row,
                "table": document.tables[table_index],
            }
            used_tables.add(table_index)
            break
    return selected


def _header_rows(table: Any, role: str) -> list[tuple[int, list[str]]]:
    rows: list[tuple[int, list[str]]] = []
    for row_index, row in enumerate(table.rows[:8]):
        cells = [_cell_text(cell) for cell in row.cells]
        rows.append((row_index, cells))
    return rows


def _contains_markers(cells: list[str], markers: tuple[str, ...]) -> bool:
    joined = "".join(_normalise(cell) for cell in cells)
    return all(_normalise(marker) in joined for marker in markers)


def _inspect_table_rows(
    role: str,
    table_info: dict[str, Any],
    empty_fields: list[dict[str, Any]],
    failures: list[dict[str, Any]],
) -> tuple[int, list[dict[str, Any]]]:
    table = table_info["table"]
    header_row = table_info["header_row"]
    rows = [(index, [_cell_text(cell) for cell in row.cells]) for index, row in enumerate(table.rows[header_row + 1 :], header_row + 1)]
    populated_rows = [(index, cells) for index, cells in rows if any(_normalise(cell) for cell in cells)]

    for row_index, cells in rows:
        if any(_normalise(cell) for cell in cells):
            continue
        _add_failure(
            failures,
            "empty_data_row",
            f"The {role} table contains an empty data row.",
            {"section": role, "row": row_index},
        )

    if role == "summary":
        _inspect_summary_fields(populated_rows, empty_fields, failures)
    else:
        header_cells = [_cell_text(cell) for cell in table.rows[header_row].cells]
        columns = _field_columns(role, header_cells)
        for row_index, cells in populated_rows:
            for field, column in columns:
                if column >= len(cells) or not _normalise(cells[column]):
                    entry = {
                        "section": role,
                        "row": row_index,
                        "column": field,
                    }
                    empty_fields.append(entry)
                    _add_failure(
                        failures,
                        "empty_field",
                        f"The {role} table has an empty required field.",
                        entry,
                    )
    return len(populated_rows), empty_fields


def _inspect_summary_fields(
    rows: list[tuple[int, list[str]]],
    empty_fields: list[dict[str, Any]],
    failures: list[dict[str, Any]],
) -> None:
    for row_index, cells in rows:
        for column in range(0, len(cells), 2):
            key = _normalise(cells[column]).rstrip("：:")
            if key not in {_normalise(value) for value in _SUMMARY_KEYS}:
                continue
            value = cells[column + 1] if column + 1 < len(cells) else ""
            if _normalise(value):
                continue
            entry = {"section": "summary", "row": row_index, "column": key}
            empty_fields.append(entry)
            _add_failure(
                failures,
                "empty_field",
                "The summary table has an empty required field.",
                entry,
            )


def _field_columns(role: str, header_cells: list[str]) -> list[tuple[str, int]]:
    columns: list[tuple[str, int]] = []
    for field, aliases in _FIELD_ALIASES[role]:
        for column, cell in enumerate(header_cells):
            normalized = _normalise(cell)
            if any(_normalise(alias) == normalized or _normalise(alias) in normalized for alias in aliases):
                columns.append((field, column))
                break
    return columns


def _cell_text(cell: Any) -> str:
    return str(cell.text or "").replace("\u00a0", " ").strip()


def _normalise(value: str) -> str:
    return re.sub(r"\s+", "", value).casefold()
