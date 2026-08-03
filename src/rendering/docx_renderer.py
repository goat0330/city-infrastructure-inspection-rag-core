"""Render Gold records and inspection predictions into a readable DOCX."""

from __future__ import annotations

from collections.abc import Mapping, Sequence
from dataclasses import asdict, is_dataclass
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt

from ..contracts.prediction import InspectionPrediction


SUMMARY_FIELDS = (
    ("bridge_name", "桥梁名称"),
    ("bridge_id", "桥梁编号"),
    ("report_date", "报告日期"),
    ("overall_score", "总体评分"),
    ("overall_grade", "总体等级"),
    ("superstructure_score", "上部结构评分"),
    ("superstructure_grade", "上部结构等级"),
    ("substructure_score", "下部结构评分"),
    ("substructure_grade", "下部结构等级"),
    ("deck_score", "桥面系评分"),
    ("deck_grade", "桥面系等级"),
    ("previous_overall_score", "上一次总体评分"),
    ("previous_overall_grade", "上一次总体等级"),
    ("trend", "病害发展趋势与具体说明"),
    ("overall_conclusion", "总体结论"),
    ("risk_points", "主要风险点"),
    ("recommendations_summary", "建议"),
)

RECOMMENDATION_COLUMNS = (
    ("index", "序号"),
    ("category", "建议类别"),
    ("content", "建议内容"),
    ("location", "病害部位"),
)

DEFECT_COLUMNS = (
    ("index", "序号"),
    ("location", "病害部位"),
    ("defect_type", "病害类型"),
    ("description", "病害描述"),
    ("is_new", "是否新增"),
    ("previous_status", "历史状态"),
    ("development", "发展"),
)

TEXT_SECTIONS = (
    ("detailed_conclusion", "详细结论"),
    ("causes", "病害成因"),
    ("treatments", "处置建议"),
    ("safety_impact", "安全影响"),
)


def _text(value: Any) -> str:
    """Convert a contract value to text without inventing content."""

    if value is None:
        return ""
    return str(value)


def _items(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, (str, bytes)):
        return [value]
    if isinstance(value, Sequence):
        return list(value)
    return [value]


def _mapping(value: Any) -> Mapping[str, Any]:
    if isinstance(value, Mapping):
        return value
    if is_dataclass(value):
        converted = asdict(value)
        if isinstance(converted, Mapping):
            return converted
    to_dict = getattr(value, "to_dict", None)
    if callable(to_dict):
        converted = to_dict()
        if isinstance(converted, Mapping):
            return converted
    return {}


def _load_source(source: Any) -> Mapping[str, Any] | list[Any]:
    if isinstance(source, Mapping):
        return source
    if isinstance(source, InspectionPrediction):
        return source.to_dict()
    if isinstance(source, Path):
        return _load_json_file(source)
    if isinstance(source, str):
        stripped = source.lstrip()
        if stripped.startswith(("{", "[")):
            loaded = json.loads(source)
            if isinstance(loaded, (Mapping, list)):
                return loaded
            raise ValueError("JSON input must be an object or a list")
        return _load_json_file(Path(source))
    converted = _mapping(source)
    if converted:
        return converted
    raise TypeError("input must be an InspectionPrediction, mapping, JSON text, or JSON path")


def _load_json_file(path: Path) -> Mapping[str, Any] | list[Any]:
    loaded = json.loads(path.read_text(encoding="utf-8-sig"))
    if isinstance(loaded, (Mapping, list)):
        return loaded
    raise ValueError("JSON input must be an object or a list")


def _record(source: Any) -> Mapping[str, Any]:
    payload = _load_source(source)
    if isinstance(payload, list):
        records = payload
        if len(records) != 1:
            raise ValueError("rendering requires exactly one inspection record")
        payload = records[0]
    if not isinstance(payload, Mapping):
        raise ValueError("inspection record must be a JSON object")

    if "records" in payload and "summary" not in payload:
        records = payload["records"]
        if not isinstance(records, list) or len(records) != 1:
            raise ValueError("Gold JSON must contain exactly one record")
        payload = records[0]
    if not isinstance(payload, Mapping):
        raise ValueError("inspection record must be a JSON object")
    return payload


def _row_value(row: Any, field: str) -> str:
    if isinstance(row, Mapping):
        return _text(row.get(field, ""))
    return _text(getattr(row, field, ""))


def _set_cell(cell: Any, value: Any) -> None:
    cell.text = _text(value)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)


def _add_table(
    document: Any,
    columns: Sequence[tuple[str, str]],
    rows: Any,
) -> Any:
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for cell, (_, label) in zip(table.rows[0].cells, columns):
        _set_cell(cell, label)

    data_rows = _items(rows)
    for row_data in data_rows:
        row = table.add_row()
        for cell, (field, _) in zip(row.cells, columns):
            _set_cell(cell, _row_value(row_data, field))
    _merge_repeated_indexes(table, len(data_rows))
    return table


def _merge_repeated_indexes(table: Any, data_row_count: int) -> None:
    """Vertically merge only adjacent non-empty equal index cells."""

    if data_row_count < 2:
        return
    values = [table.rows[row_index + 1].cells[0].text.strip() for row_index in range(data_row_count)]
    start = 0
    while start < data_row_count:
        value = values[start]
        end = start
        while value and end + 1 < data_row_count and values[end + 1] == value:
            end += 1
        if end > start:
            merged = table.cell(start + 1, 0).merge(table.cell(end + 1, 0))
            _set_cell(merged, value)
            merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        start = end + 1


def _add_text_section(document: Any, title: str, values: Any) -> None:
    document.add_heading(title, level=1)
    for value in _items(values):
        document.add_paragraph(_text(value))


def _render_record(record: Mapping[str, Any], output_path: Path | str) -> Path:
    document = Document()
    document.add_heading("信息提取报告", level=0)

    sample_id = _text(record.get("sample_id", ""))
    if sample_id:
        document.add_paragraph(sample_id)

    document.add_heading("概要表", level=1)
    summary = _mapping(record.get("summary"))
    summary_rows: list[tuple[str, str]] = []
    if sample_id:
        summary_rows.append(("样本编号", sample_id))
    summary_rows.extend((label, _text(summary.get(field, ""))) for field, label in SUMMARY_FIELDS)
    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    for cell, value in zip(summary_table.rows[0].cells, ("字段", "内容")):
        _set_cell(cell, value)
    for label, value in summary_rows:
        row = summary_table.add_row()
        _set_cell(row.cells[0], label)
        _set_cell(row.cells[1], value)

    for field, title in TEXT_SECTIONS[:1]:
        _add_text_section(document, title, record.get(field))

    document.add_heading("建议明细", level=1)
    _add_table(document, RECOMMENDATION_COLUMNS, record.get("recommendations"))

    document.add_heading("病害列表", level=1)
    _add_table(document, DEFECT_COLUMNS, record.get("defects"))

    for field, title in TEXT_SECTIONS[1:]:
        _add_text_section(document, title, record.get(field))

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


def render_report(source: Any, output_path: Path | str) -> Path:
    """Render one Gold record or :class:`InspectionPrediction` to ``output_path``."""

    return _render_record(_record(source), output_path)


def render_prediction(prediction: Any, output_path: Path | str) -> Path:
    return render_report(prediction, output_path)


def render_gold(record: Any, output_path: Path | str) -> Path:
    return render_report(record, output_path)


render_docx = render_report
render_prediction_to_docx = render_prediction
render_inspection_prediction = render_prediction
render_gold_to_docx = render_gold
