"""Render structured inspection predictions into the production DOCX template."""

from __future__ import annotations

from copy import deepcopy
import json
from pathlib import Path
import re
from typing import Any, Mapping, Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.table import _Row
from docx.text.paragraph import Paragraph

from .style_spec import (
    CAUSE_NUMBER_FORMAT,
    DEFECT_CENTERED_COLUMNS,
    DEFECT_LEFT_ALIGNED_COLUMNS,
    RECOMMENDATION_CENTERED_COLUMNS,
    RECOMMENDATION_LEFT_ALIGNED_COLUMNS,
    TREATMENT_NUMBER_FORMAT,
)
from .submission_document import SubmissionDocument, build_submission_document


_REPO_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_TEMPLATE_PATH = _REPO_ROOT / "assets" / "templates" / "information_extraction_v1.docx"
DEFAULT_FIELDS_PATH = _REPO_ROOT / "assets" / "templates" / "template_fields.json"
_PLACEHOLDER_RE = re.compile(r"\{\{[^{}]+\}\}")
_REFERENCE_RED = RGBColor(0xE5, 0x4C, 0x5E)
_OUTPUT_INK = RGBColor(0x1E, 0x2A, 0x36)
_REFERENCE_DISPLAY_TEXT = {
    "1、简要信息": "1、简要信息（20分）",
    "从报告中提取桥梁定检的概要结论与关键指标，输出内容包括：": "从定检报告中提取概要结论与关键指标，输出内容包括：",
    "2、详细信息": "2、详细信息（80分）",
    "在简要信息基础上，进一步提取以下结构化明细：": "在简要信息基础上，进一步提取以下结构化明细：",
    "（1）详细结论": "（1）详细结论（15分）",
    "（2）建议明细": "（2）建议明细（20分）",
    "病害列表": "病害列表（30分）",
    "病害成因": "病害成因（5分）：",
    "处置建议": "处置建议（5分）：",
    "安全影响": "安全影响（5分）：",
}
_REFERENCE_PARAGRAPHS = set(_REFERENCE_DISPLAY_TEXT.values())


def _load_contract(path: str | Path | None) -> Mapping[str, Any]:
    source = Path(path) if path is not None else DEFAULT_FIELDS_PATH
    if not source.is_file():
        raise FileNotFoundError(f"template field contract not found: {source}")
    value = json.loads(source.read_text(encoding="utf-8-sig"))
    if not isinstance(value, Mapping):
        raise ValueError("template field contract must be a JSON object")
    return value


def _remove_run(run: Any) -> None:
    parent = run._element.getparent()
    if parent is not None:
        parent.remove(run._element)


def _set_paragraph_text(
    paragraph: Paragraph,
    value: object,
    *,
    color: RGBColor | None = None,
) -> None:
    """Replace the whole slot text while preserving the first run formatting."""

    text = "" if value is None else str(value)
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            _remove_run(run)
    else:
        paragraph.add_run(text)
    if color is not None:
        for run in paragraph.runs:
            run.font.color.rgb = color


def _color_paragraph(paragraph: Paragraph, color: RGBColor) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = color


def _set_cell_text(
    cell: Any,
    value: object,
    *,
    align: int | None = None,
    color: RGBColor | None = _OUTPUT_INK,
) -> None:
    while len(cell.paragraphs) > 1:
        paragraph = cell.paragraphs[-1]._element
        paragraph.getparent().remove(paragraph)
    paragraph = cell.paragraphs[0]
    _set_paragraph_text(paragraph, value, color=color)
    paragraph.paragraph_format.space_after = 0
    if align is not None:
        paragraph.alignment = align
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _find_paragraph(document: Document, text: str) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == text]
    if len(matches) != 1:
        raise ValueError(f"expected one paragraph {text!r}, found {len(matches)}")
    return matches[0]


def _replace_exact_placeholder(
    document: Document,
    placeholder: str,
    value: object,
    *,
    color: RGBColor | None = _OUTPUT_INK,
) -> None:
    matches: list[Paragraph] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == placeholder:
            matches.append(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip() == placeholder:
                        matches.append(paragraph)
    if len(matches) != 1:
        raise ValueError(f"placeholder {placeholder!r} expected once, found {len(matches)}")
    _set_paragraph_text(matches[0], value, color=color)


def _remove_row(table: Any, row: _Row) -> None:
    table._tbl.remove(row._tr)


def _clone_row_before(table: Any, prototype: _Row) -> _Row:
    tr = deepcopy(prototype._tr)
    prototype._tr.addprevious(tr)
    return _Row(tr, table)


def _ensure_cant_split(row: _Row) -> None:
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))


def _render_recommendations(document: Document, contract: Mapping[str, Any], values: Sequence[Mapping[str, str]]) -> None:
    spec = contract["repeaters"]["recommendations"]
    table = document.tables[int(spec["table_index"])]
    prototype = table.rows[int(spec["prototype_row"])]
    for number, item in enumerate(values, 1):
        row = _clone_row_before(table, prototype)
        _ensure_cant_split(row)
        payload = {
            0: str(number),
            1: item.get("category", ""),
            2: item.get("content", ""),
            3: item.get("location", ""),
        }
        for index, cell in enumerate(row.cells):
            align = WD_ALIGN_PARAGRAPH.LEFT if index in RECOMMENDATION_LEFT_ALIGNED_COLUMNS else WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_text(cell, payload[index], align=align)
    _remove_row(table, prototype)


def _normalized_defect_indexes(values: Sequence[Mapping[str, str]]) -> list[str]:
    result: list[str] = []
    previous_key: str | None = None
    display = 0
    for position, item in enumerate(values):
        raw = str(item.get("index", "")).strip()
        key = raw or previous_key or f"__row_{position}"
        if previous_key is None or key != previous_key:
            display += 1
        result.append(str(display))
        previous_key = key
    return result


def _merge_consecutive_indexes(table: Any) -> None:
    if len(table.rows) <= 2:
        return
    values = [row.cells[0].text.strip() for row in table.rows[1:]]
    start = 0
    while start < len(values):
        end = start
        while end + 1 < len(values) and values[end + 1] == values[start]:
            end += 1
        if end > start and values[start]:
            merged = table.cell(start + 1, 0).merge(table.cell(end + 1, 0))
            _set_cell_text(merged, values[start], align=WD_ALIGN_PARAGRAPH.CENTER)
        start = end + 1


def _render_defects(document: Document, contract: Mapping[str, Any], values: Sequence[Mapping[str, str]]) -> None:
    spec = contract["repeaters"]["defects"]
    table = document.tables[int(spec["table_index"])]
    prototype = table.rows[int(spec["prototype_row"])]
    indexes = _normalized_defect_indexes(values)
    for number, item in zip(indexes, values):
        row = _clone_row_before(table, prototype)
        _ensure_cant_split(row)
        payload = {
            0: number,
            1: item.get("location", ""),
            2: item.get("type", ""),
            3: item.get("description", ""),
            4: item.get("is_new", ""),
            5: item.get("previous_status", ""),
            6: item.get("development_degree", ""),
        }
        for index, cell in enumerate(row.cells):
            align = WD_ALIGN_PARAGRAPH.LEFT if index in DEFECT_LEFT_ALIGNED_COLUMNS else WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_text(cell, payload[index], align=align)
    _remove_row(table, prototype)
    _merge_consecutive_indexes(table)


def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _copy_run_properties(source: Any, target: Any) -> None:
    source_properties = deepcopy(source._r.rPr)
    existing = target._r.rPr
    if existing is not None:
        target._r.remove(existing)
    if source_properties is not None:
        target._r.insert(0, source_properties)


def _color_text_fragments(paragraph: Paragraph, patterns: Sequence[str]) -> None:
    """Color only fixed cue phrases while leaving extracted facts dark."""

    text = paragraph.text
    if not text or not paragraph.runs:
        return
    matches: list[tuple[int, int]] = []
    for pattern in patterns:
        matches.extend((match.start(), match.end()) for match in re.finditer(pattern, text))
    if not matches:
        return
    matches.sort()
    selected: list[tuple[int, int]] = []
    for start, end in matches:
        if selected and start < selected[-1][1]:
            selected[-1] = (selected[-1][0], max(selected[-1][1], end))
        else:
            selected.append((start, end))
    segments: list[tuple[str, bool]] = []
    cursor = 0
    for start, end in selected:
        if start > cursor:
            segments.append((text[cursor:start], False))
        segments.append((text[start:end], True))
        cursor = end
    if cursor < len(text):
        segments.append((text[cursor:], False))

    base = paragraph.runs[0]
    base_properties = deepcopy(base._r.rPr)
    for run in paragraph.runs[1:]:
        _remove_run(run)
    base.text = segments[0][0]
    if segments[0][1]:
        base.font.color.rgb = _REFERENCE_RED
    else:
        base.font.color.rgb = _OUTPUT_INK
    for segment, is_reference in segments[1:]:
        run = paragraph.add_run(segment)
        if base_properties is not None:
            _copy_run_properties(base, run)
        run.font.color.rgb = _REFERENCE_RED if is_reference else _OUTPUT_INK


def _clone_paragraph_before(prototype: Paragraph) -> Paragraph:
    element = deepcopy(prototype._p)
    prototype._p.addprevious(element)
    return Paragraph(element, prototype._parent)


def _render_paragraph_repeater(
    document: Document,
    contract: Mapping[str, Any],
    name: str,
    values: Sequence[str],
) -> None:
    spec = contract["repeaters"][name]
    heading = _find_paragraph(document, str(spec["anchor_heading"]))
    paragraphs = document.paragraphs
    heading_index = next(i for i, item in enumerate(paragraphs) if item._p is heading._p)
    offset = int(spec.get("prototype_paragraph_offset", 1))
    if heading_index + offset >= len(paragraphs):
        raise ValueError(f"prototype paragraph missing after {heading.text!r}")
    prototype = paragraphs[heading_index + offset]
    numbering = spec.get("numbering", {})
    mode = str(numbering.get("mode", "none"))
    fmt = str(numbering.get("format", "{n}"))
    for number, value in enumerate(values, int(numbering.get("start", 1))):
        paragraph = _clone_paragraph_before(prototype)
        prefix = fmt.format(n=number) if mode == "renderer_prefix" else ""
        _set_paragraph_text(paragraph, f"{prefix}{value}", color=_OUTPUT_INK)
    _delete_paragraph(prototype)


def _apply_official_reference_colors(document: Document) -> None:
    """Match the official example: fixed cues red, extracted text dark."""

    for paragraph in document.paragraphs:
        stored_text = paragraph.text.strip()
        display_text = _REFERENCE_DISPLAY_TEXT.get(stored_text)
        if display_text is not None:
            _set_paragraph_text(paragraph, display_text, color=_REFERENCE_RED)
        elif stored_text in _REFERENCE_PARAGRAPHS:
            _color_paragraph(paragraph, _REFERENCE_RED)
            continue
        if paragraph.text.startswith("经综合评定"):
            _color_text_fragments(
                paragraph,
                (
                    r"经综合评定",
                    r"(?:该|本)[^，。；]{0,20}总体技术状况评分",
                    r"总体技术状况等级为",
                    r"上部结构评分",
                    r"下部结构评分",
                    r"桥面系评分",
                ),
            )
        elif paragraph.text.startswith("目前，"):
            _color_text_fragments(paragraph, (r"目前，", r"上部结构", r"下部结构", r"桥面系"))
        elif paragraph.text.startswith("综上，"):
            _color_text_fragments(paragraph, (r"综上，", r"(?<=；)建议"))


def _all_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    values.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join(values)


def _validate_rendered(document: Document) -> None:
    remaining = sorted(set(_PLACEHOLDER_RE.findall(_all_text(document))))
    if remaining:
        raise ValueError(f"rendered document contains unresolved placeholders: {remaining[:10]}")
    if len(document.tables) != 3:
        raise ValueError(f"production template must contain exactly three tables, found {len(document.tables)}")
    if len(document.tables[0].columns) != 3:
        raise ValueError("summary table must contain three columns")
    if len(document.tables[1].columns) != 4:
        raise ValueError("recommendation table must contain four columns")
    if len(document.tables[2].columns) != 7:
        raise ValueError("defect table must contain seven columns")


def render_template_report(
    source: object,
    output_path: str | Path,
    *,
    template_path: str | Path | None = None,
    fields_path: str | Path | None = None,
    facility_context: object = None,
    field_states: object = None,
) -> Path:
    """Render one prediction using the frozen production DOCX template."""

    submission = (
        source
        if isinstance(source, SubmissionDocument)
        else build_submission_document(
            source,
            facility_context=facility_context,
            field_states=field_states,
        )
    )
    template = Path(template_path) if template_path is not None else DEFAULT_TEMPLATE_PATH
    if not template.is_file():
        raise FileNotFoundError(f"production template not found: {template}")
    contract = _load_contract(fields_path)
    document = Document(template)

    for field, value in submission.scalars.items():
        _replace_exact_placeholder(document, "{{" + field + "}}", value)
    for field in (
        "score_and_grade",
        "history_and_defects",
        "current_structure_state",
        "comprehensive_judgement",
    ):
        _replace_exact_placeholder(document, "{{" + field + "}}", getattr(submission, field))

    _render_recommendations(document, contract, submission.recommendations)
    _render_defects(document, contract, submission.defects)
    _render_paragraph_repeater(document, contract, "causes", submission.causes)
    _render_paragraph_repeater(document, contract, "treatments", submission.treatments)
    _render_paragraph_repeater(document, contract, "safety_impacts", submission.safety_impacts)
    _apply_official_reference_colors(document)
    _validate_rendered(document)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output
