import json
import tempfile
import unittest
from pathlib import Path
from subprocess import CompletedProcess

from docx import Document

from src.conversion import convert_directory, convert_docx_directory


class FakeSoffice:
    def __init__(self, failing_names: set[str] | None = None) -> None:
        self.failing_names = failing_names or set()
        self.calls: list[list[str]] = []

    def __call__(self, command: list[str]) -> CompletedProcess[str]:
        self.calls.append(command)
        source = Path(command[-1])
        if source.name in self.failing_names:
            return CompletedProcess(command, 1, stdout="", stderr="synthetic failure")

        outdir = Path(command[command.index("--outdir") + 1])
        output = outdir / f"{source.stem}.docx"
        Document().save(output)
        return CompletedProcess(command, 0, stdout="converted", stderr="")


class FakeDocxSoffice:
    def __init__(self) -> None:
        self.calls: list[list[str]] = []

    def __call__(self, command: list[str]) -> CompletedProcess[str]:
        self.calls.append(command)
        source = Path(command[-1])
        outdir = Path(command[command.index("--outdir") + 1])
        output = outdir / f"{source.stem}.doc"
        output.write_bytes(b"synthetic legacy doc")
        return CompletedProcess(command, 0, stdout="converted", stderr="")


class ConversionTests(unittest.TestCase):
    def make_workspace(self) -> tuple[tempfile.TemporaryDirectory[str], Path, Path, Path]:
        holder = tempfile.TemporaryDirectory(
            prefix=".conversion-test-", dir=Path(__file__).resolve().parent
        )
        root = Path(holder.name)
        input_dir = root / "input"
        output_dir = root / "output"
        state_path = root / "state.json"
        input_dir.mkdir()
        return holder, input_dir, output_dir, state_path

    def test_failure_isolated_and_state_has_required_fields(self) -> None:
        holder, input_dir, output_dir, state_path = self.make_workspace()
        with holder:
            (input_dir / "ok.doc").write_bytes(b"synthetic legacy document")
            (input_dir / "broken.doc").write_bytes(b"synthetic broken document")
            fake = FakeSoffice({"broken.doc"})

            result = convert_directory(
                input_dir, output_dir, state_path, "fake-soffice", runner=fake
            )

            self.assertEqual(result.counts, {"success": 1, "skipped": 0, "failed": 1})
            records = {Path(record["source"]).name: record for record in result.records}
            self.assertEqual(records["ok.doc"]["status"], "success")
            self.assertEqual(records["broken.doc"]["status"], "failed")
            self.assertTrue(Path(records["ok.doc"]["target"]).is_file())
            Document(records["ok.doc"]["target"])
            self.assertIn("synthetic failure", records["broken.doc"]["error"])
            required = {
                "source",
                "target",
                "status",
                "duration",
                "duration_ms",
                "source_size",
                "source_sha256",
                "source_mtime_ns",
                "target_size",
                "target_mtime_ns",
                "target_is_usable",
                "error",
            }
            for record in records.values():
                self.assertEqual(set(record), required)

            state = json.loads(state_path.read_text(encoding="utf-8"))
            self.assertEqual(len(state["records"]), 2)

    def test_success_is_skipped_until_source_changes(self) -> None:
        holder, input_dir, output_dir, state_path = self.make_workspace()
        with holder:
            source = input_dir / "report.doc"
            source.write_bytes(b"first synthetic document")
            first_fake = FakeSoffice()
            first = convert_directory(
                input_dir, output_dir, state_path, "fake-soffice", runner=first_fake
            )
            self.assertEqual(first.counts["success"], 1)

            second_fake = FakeSoffice()
            second = convert_directory(
                input_dir, output_dir, state_path, "fake-soffice", runner=second_fake
            )
            self.assertEqual(second.records[0]["status"], "skipped")
            self.assertEqual(second_fake.calls, [])

            source.write_bytes(b"a changed synthetic document")
            third_fake = FakeSoffice()
            third = convert_directory(
                input_dir, output_dir, state_path, "fake-soffice", runner=third_fake
            )
            self.assertEqual(third.records[0]["status"], "success")
            self.assertEqual(len(third_fake.calls), 1)

    def test_each_conversion_uses_a_unique_libreoffice_profile(self) -> None:
        holder, input_dir, output_dir, state_path = self.make_workspace()
        with holder:
            for name in ("one.doc", "two.doc"):
                (input_dir / name).write_bytes(b"synthetic legacy document")
            fake = FakeSoffice()

            convert_directory(
                input_dir, output_dir, state_path, "fake-soffice", runner=fake
            )

            profiles = [
                argument.split("=", 1)[1]
                for call in fake.calls
                for argument in call
                if argument.startswith("-env:UserInstallation=")
            ]
            self.assertEqual(len(profiles), 2)
            self.assertEqual(len(set(profiles)), 2)

    def test_same_size_source_change_is_not_skipped(self) -> None:
        holder, input_dir, output_dir, state_path = self.make_workspace()
        with holder:
            source = input_dir / "report.doc"
            source.write_bytes(b"AAAA")
            first_fake = FakeSoffice()
            convert_directory(input_dir, output_dir, state_path, "fake-soffice", runner=first_fake)

            source.write_bytes(b"BBBB")
            second_fake = FakeSoffice()
            second = convert_directory(input_dir, output_dir, state_path, "fake-soffice", runner=second_fake)
            self.assertEqual(second.records[0]["status"], "success")
            self.assertEqual(len(second_fake.calls), 1)

    def test_timeout_is_recorded_as_failure(self) -> None:
        holder, input_dir, output_dir, state_path = self.make_workspace()
        with holder:
            (input_dir / "slow.doc").write_bytes(b"synthetic")

            def timeout_runner(command: list[str]):
                raise TimeoutError("synthetic timeout")

            result = convert_directory(
                input_dir, output_dir, state_path, "fake-soffice", runner=timeout_runner
            )
            self.assertEqual(result.counts["failed"], 1)
            self.assertIn("TimeoutError", result.records[0]["error"])
            self.assertFalse(result.records[0]["target_is_usable"])

    def test_docx_conversion_uses_public_temp_paths(self) -> None:
        holder = tempfile.TemporaryDirectory(prefix="conversion-docx-test-")
        root = Path(holder.name)
        try:
            input_dir = root / "input"
            output_dir = root / "output"
            input_dir.mkdir()
            for name in ("one.docx", "two.docx"):
                Document().save(input_dir / name)

            fake = FakeDocxSoffice()
            result = convert_docx_directory(
                input_dir,
                output_dir,
                soffice_path="fake-soffice",
                runner=fake,
                engine="libreoffice",
            )

            self.assertEqual(result.counts, {"success": 2, "skipped": 0, "failed": 0})
            self.assertEqual(len(fake.calls), 2)
            for call in fake.calls:
                profile_argument = next(
                    value for value in call if value.startswith("-env:UserInstallation=")
                )
                self.assertNotIn("/.docx-to-doc-", profile_argument.replace("\\", "/"))
                self.assertNotIn("/.conversion-", profile_argument.replace("\\", "/"))
        finally:
            holder.cleanup()


if __name__ == "__main__":
    unittest.main()
