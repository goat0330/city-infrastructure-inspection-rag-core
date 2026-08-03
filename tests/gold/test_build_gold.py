from __future__ import annotations

import json
from pathlib import Path
import sys
import subprocess
import tempfile
import unittest

from docx import Document

REPO_ROOT = Path(__file__).parents[2]
sys.path.insert(0, str(REPO_ROOT))

from scripts.build_gold import build_gold  # noqa: E402


def write_label(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    summary = document.add_table(rows=3, cols=2)
    summary.cell(0, 0).text = "项目"
    summary.cell(0, 1).text = "值"
    summary.cell(1, 0).text = "桥梁名称"
    summary.cell(1, 1).text = "示例桥"
    summary.cell(2, 0).text = "总体评分"
    summary.cell(2, 1).text = "82.00"
    recommendations = document.add_table(rows=2, cols=4)
    for cell, value in zip(recommendations.rows[0].cells, ("序号", "建议类别", "建议内容", "部位")):
        cell.text = value
    for cell, value in zip(recommendations.rows[1].cells, ("1", "尽快维修", "修复桥面", "桥面")):
        cell.text = value
    defects = document.add_table(rows=2, cols=7)
    for cell, value in zip(defects.rows[0].cells, ("序号", "病害部位", "病害类型", "病害描述", "是否新增", "历史状态", "发展")):
        cell.text = value
    for cell, value in zip(defects.rows[1].cells, ("1", "桥面", "裂缝", "一条裂缝", "是", "无", "无")):
        cell.text = value
    document.add_paragraph("（1）详细结论")
    document.add_paragraph("结构状况良好")
    document.add_paragraph("（2）建议明细")
    document.add_paragraph("病害成因")
    document.add_paragraph("雨水侵蚀")
    document.add_paragraph("处置建议")
    document.add_paragraph("及时修复")
    document.add_paragraph("安全影响")
    document.add_paragraph("影响行车安全")
    document.save(path)


class GoldBuildTests(unittest.TestCase):
    def test_gold_fields_failures_and_repeated_output_are_stable(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            labels = root / "labels"
            reports = root / "reports"
            write_label(labels / "2013年/示例桥-无对比年度的信息提取报告.docx")
            legacy = labels / "2013年/旧标签-无对比年度的信息提取报告.doc"
            legacy.write_bytes(b"legacy")
            reports.mkdir()
            (reports / "示例桥.docx").write_bytes(b"report")

            first = build_gold(labels, reports, root / "out-1")
            second = build_gold(labels, reports, root / "out-2")
            self.assertEqual(first, second)
            self.assertEqual(first["statistics"]["label_count"], 2)
            self.assertEqual(first["statistics"]["record_count"], 1)
            self.assertEqual(first["statistics"]["failed_count"], 1)
            self.assertEqual(first["statistics"]["quality_flag_count"], 0)
            record = first["records"][0]
            self.assertEqual(record["summary"]["bridge_name"], "示例桥")
            self.assertEqual(record["recommendations"][0]["content"], "修复桥面")
            self.assertEqual(record["defects"][0]["defect_type"], "裂缝")
            self.assertEqual(record["causes"], ["雨水侵蚀"])
            self.assertEqual(record["treatments"], ["及时修复"])
            self.assertEqual(record["safety_impact"], ["影响行车安全"])
            self.assertFalse(record["provenance"]["raw_report_included"])
            self.assertEqual(record["provenance"]["source_report_relative_path"], "示例桥.docx")
            self.assertEqual(first["failed"][0]["error_code"], "legacy_doc_unsupported")
            self.assertEqual(
                (root / "out-1/gold.json").read_bytes(),
                (root / "out-2/gold.json").read_bytes(),
            )
            result = subprocess.run(
                [
                    sys.executable,
                    str(REPO_ROOT / "scripts/build_gold.py"),
                    "--labels-dir",
                    str(labels),
                    "--reports-dir",
                    str(reports),
                    "--output-dir",
                    str(root / "out-cli"),
                ],
                cwd=REPO_ROOT,
                check=True,
                capture_output=True,
                text=True,
            )
            self.assertEqual(json.loads(result.stdout)["record_count"], 1)
            self.assertTrue((root / "out-cli/gold.json").is_file())

    def test_parser_locates_tables_by_headers_not_fixed_positions(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            labels = root / "labels"
            reports = root / "reports"
            path = labels / "2012年/额外说明桥-无对比年度的信息提取报告.docx"
            path.parent.mkdir(parents=True, exist_ok=True)
            document = Document()
            extra = document.add_table(rows=1, cols=1)
            extra.cell(0, 0).text = "说明表"
            document.add_paragraph("标签说明")
            summary = document.add_table(rows=3, cols=3)
            for cell, value in zip(summary.rows[0].cells, ("字段", "内容", "说明")):
                cell.text = value
            for cell, value in zip(summary.rows[1].cells, ("桥梁名称", "额外说明桥", "全称")):
                cell.text = value
            for cell, value in zip(summary.rows[2].cells, ("总体评分", "90.0", "评分")):
                cell.text = value
            recs = document.add_table(rows=2, cols=4)
            for cell, value in zip(recs.rows[0].cells, ("序号", "建议类别", "建议内容", "病害部位")):
                cell.text = value
            for cell, value in zip(recs.rows[1].cells, ("1", "尽快维修", "修复", "桥面")):
                cell.text = value
            defects = document.add_table(rows=2, cols=7)
            for cell, value in zip(defects.rows[0].cells, ("序号", "病害部位", "病害类型", "病害描述", "是否新增", "上一次定检状态", "发展程度")):
                cell.text = value
            for cell, value in zip(defects.rows[1].cells, ("1", "桥面", "裂缝", "裂缝一条", "否", "无", "无")):
                cell.text = value
            document.save(path)
            reports.mkdir()
            (reports / "额外说明桥.doc").write_bytes(b"report")
            payload = build_gold(labels, reports, root / "out")
            record = payload["records"][0]
            self.assertEqual(record["summary"]["bridge_name"], "额外说明桥")
            self.assertEqual(record["recommendations"][0]["content"], "修复")
            self.assertEqual(record["defects"][0]["description"], "裂缝一条")


if __name__ == "__main__":
    unittest.main()
