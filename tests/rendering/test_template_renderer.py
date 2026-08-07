from __future__ import annotations

from pathlib import Path
import re

from docx import Document

from src.contracts import BridgeSummary, DefectObservation, InspectionPrediction, Recommendation
from src.rendering import build_submission_document, render_template_report


TEMPLATE = Path("assets/templates/information_extraction_v1.docx")
FIELDS = Path("assets/templates/template_fields.json")


def _prediction() -> InspectionPrediction:
    return InspectionPrediction(
        sample_id="synthetic",
        source_file="synthetic.docx",
        summary=BridgeSummary(
            bridge_name="测试桥",
            report_date="2026年8月",
            overall_score="88.00",
            overall_grade="B级",
            superstructure_score="90.00",
            superstructure_grade="A级",
            substructure_score="86.00",
            substructure_grade="B级",
            deck_score="80.00",
            deck_grade="C级",
            previous_overall_score="无",
            previous_overall_grade="无",
            trend="无",
            overall_conclusion="总体状况良好。",
            risk_points="支座开裂影响耐久性。",
            recommendations_summary="1条尽快维修、1条预防性养护",
        ),
        detailed_conclusion=("评分段。", "历史段。", "状态段。", "综合段。"),
        recommendations=(
            Recommendation(index="8", category="尽快维修", content="修复支座。", location="支座"),
            Recommendation(index="20", category="预防性养护", content="加强巡检。", location="桥梁"),
        ),
        defects=(
            DefectObservation(index="3", location="支座", defect_type="开裂", description="第一处", is_new="否", previous_status="无", development="无"),
            DefectObservation(index="3", location="支座", defect_type="开裂", description="第二处", is_new="否", previous_status="无", development="无"),
            DefectObservation(index="9", location="桥面", defect_type="破损", description="第三处", is_new="否", previous_status="无", development="无"),
        ),
        causes=("材料老化。",),
        treatments=("修复支座。", "加强巡检。"),
        safety_impact=("可能影响耐久性。",),
    )


def _all_text(document: Document) -> str:
    values = [p.text for p in document.paragraphs]
    values.extend(c.text for t in document.tables for r in t.rows for c in r.cells)
    return "\n".join(values)


def test_production_template_is_dynamic_prototype() -> None:
    document = Document(TEMPLATE)
    assert len(document.tables) == 3
    assert [len(table.rows) for table in document.tables] == [17, 2, 2]
    text = _all_text(document)
    assert "{{recommendation.content}}" in text
    assert "{{defect.description}}" in text
    assert "{{cause.text}}" in text
    assert not re.search(r"_[1-9]\}\}", text)


def test_renders_dynamic_rows_and_no_placeholders(tmp_path: Path) -> None:
    output = render_template_report(
        _prediction(),
        tmp_path / "rendered.docx",
        template_path=TEMPLATE,
        fields_path=FIELDS,
    )
    document = Document(output)
    text = _all_text(document)
    assert "{{" not in text
    assert "未提取到" not in text
    assert "1、简要信息（20分）" in text
    assert "2、详细信息（80分）" in text
    assert "病害成因（5分）：" in text
    assert document.tables[0].rows[1].cells[1].text == "测试桥"
    assert len(document.tables[1].rows) == 3
    assert [r.cells[0].text for r in document.tables[1].rows[1:]] == ["1", "2"]
    assert len(document.tables[2].rows) == 4
    assert document.tables[2].rows[1].cells[0].text == "1"
    assert document.tables[2].rows[3].cells[0].text == "2"
    assert "（1）材料老化。" in text
    assert "（2）加强巡检。" in text

    heading = next(p for p in document.paragraphs if p.text == "1、简要信息（20分）")
    assert str(heading.runs[0].font.color.rgb) == "E54C5E"
    score = next(p for p in document.paragraphs if "评分段。" in p.text)
    assert str(score.runs[0].font.color.rgb) == "1E2A36"


def test_empty_dynamic_arrays_leave_headers_only(tmp_path: Path) -> None:
    prediction = InspectionPrediction(summary=BridgeSummary(bridge_name="空样本"))
    output = render_template_report(
        prediction,
        tmp_path / "empty.docx",
        template_path=TEMPLATE,
        fields_path=FIELDS,
    )
    document = Document(output)
    assert len(document.tables[1].rows) == 1
    assert len(document.tables[2].rows) == 1
    assert "{{" not in _all_text(document)


def test_submission_mapping_uses_existing_contract_names() -> None:
    submission = build_submission_document(_prediction())
    assert submission.scalars["deck_system_score"] == "80.00"
    assert submission.scalars["major_risks"] == "支座开裂影响耐久性。"
    assert submission.scalars["report_title"] == "测试桥·信息提取报告"


def test_facility_context_replaces_bridge_fallback_in_generated_narrative() -> None:
    record = {
        "summary": {
            "bridge_name": "通道-1",
            "overall_score": "88.00",
            "overall_grade": "B级",
        },
        "facility_context": {"facility_noun": "人行通道"},
        "recommendations": [],
        "defects": [],
    }

    submission = build_submission_document(record)

    assert "该人行通道" in submission.score_and_grade
    assert "该桥" not in submission.score_and_grade
    assert "桥梁" not in submission.score_and_grade

    fallback = build_submission_document({
        "summary": record["summary"],
        "recommendations": [],
        "defects": [],
    })
    assert "该设施" in fallback.score_and_grade


def test_field_states_keep_none_distinct_from_not_extracted(tmp_path: Path) -> None:
    record = {
        "summary": {
            "bridge_name": "测试设施",
            "report_date": "",
            "overall_score": "",
            "overall_grade": "无",
            "previous_overall_score": "无",
        },
        "field_states": {
            "summary": {
                "report_date": "not_extracted",
                "overall_score": "explicit_none",
                "overall_grade": "not_applicable",
            }
        },
        "recommendations": [],
        "defects": [],
    }

    scalars = build_submission_document(record).scalars

    assert scalars["report_date"] == "无"
    assert scalars["overall_score"] == "无"
    assert scalars["overall_grade"] == "无"
    assert scalars["previous_overall_score"] == "无"

    output = render_template_report(
        record,
        tmp_path / "field-states.docx",
        template_path=TEMPLATE,
        fields_path=FIELDS,
    )
    document = Document(output)
    assert document.tables[0].rows[3].cells[1].text == "无"
    assert document.tables[0].rows[3].cells[1].text != "未提取到"


def test_summary_and_narrative_fields_keep_separate_sources() -> None:
    record = {
        "summary": {
            "bridge_name": "测试设施",
            "overall_conclusion": "概要短结论。",
        },
        "narrative": {
            "detailed_conclusion": ("详细段一。", "详细段二。", "详细段三。", "详细段四。")
        },
        "recommendations": [],
        "defects": [],
    }

    submission = build_submission_document(record)

    assert submission.scalars["overall_conclusion"] == "概要短结论。"
    assert submission.score_and_grade == "详细段一。"
    assert submission.comprehensive_judgement == "详细段四。"


def test_recommendation_summary_has_three_categories_and_preserves_source() -> None:
    derived = build_submission_document(
        {
            "summary": {"bridge_name": "测试设施"},
            "recommendations": [{"category": "尽快维修"}],
            "defects": [],
        }
    )
    assert derived.scalars["recommendations_summary"] == "0条立即处置、1条尽快维修、0条预防性养护"

    source = build_submission_document(
        {
            "summary": {
                "bridge_name": "测试设施",
                "recommendations_summary": "9条立即处置、0条尽快维修、0条预防性养护",
            },
            "recommendations": [{"category": "尽快维修"}],
            "defects": [],
        }
    )
    assert source.scalars["recommendations_summary"] == "0条立即处置、1条尽快维修、0条预防性养护"


def test_facility_render_has_no_unresolved_placeholders(tmp_path: Path) -> None:
    output = render_template_report(
        {
            "summary": {"bridge_name": "人行通道"},
            "facility_context": {"subject": "该人行通道"},
            "recommendations": [],
            "defects": [],
        },
        tmp_path / "facility.docx",
        template_path=TEMPLATE,
        fields_path=FIELDS,
    )

    assert "{{" not in _all_text(Document(output))


def test_structured_aliases_land_in_exact_official_scalar_rows(tmp_path: Path) -> None:
    record = {
        "summary": {
            "bridge_name": "官方院子桥式通道",
            "report_date": "2019年11月20日",
            "overall_score": "94.80",
            "overall_grade": "A级",
            "superstructure_score": "93.10",
            "superstructure_grade": "A级",
            "substructure_score": "92.20",
            "substructure_grade": "A级",
            "deck_score": "91.30",
            "deck_grade": "A级",
            "previous_overall_score": "86.05",
            "previous_overall_grade": "B级",
            "trend": "桥面系：新增局部破损",
            "risk_points": "桥面局部破损影响耐久性。",
        },
        "recommendations": [],
        "defects": [],
    }

    submission = build_submission_document(record)
    assert submission.scalars["deck_system_score"] == "91.30"
    assert submission.scalars["deck_system_grade"] == "A级"
    assert submission.scalars["defect_development_trend"] == "桥面系：新增局部破损"
    assert submission.scalars["major_risks"] == "桥面局部破损影响耐久性。"

    output = render_template_report(
        record,
        tmp_path / "structured-aliases.docx",
        template_path=TEMPLATE,
        fields_path=FIELDS,
    )
    document = Document(output)
    expected = {
        1: "官方院子桥式通道",
        2: "2019年11月20日",
        3: "94.80",
        4: "A级",
        5: "93.10",
        6: "A级",
        7: "92.20",
        8: "A级",
        9: "91.30",
        10: "A级",
        11: "86.05",
        12: "B级",
        13: "桥面系：新增局部破损",
        15: "桥面局部破损影响耐久性。",
    }
    for row_index, value in expected.items():
        assert document.tables[0].rows[row_index].cells[1].text == value
