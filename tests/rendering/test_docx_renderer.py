from __future__ import annotations

from pathlib import Path

from docx import Document

from src.contracts import BridgeSummary, DefectObservation, InspectionPrediction, Recommendation
from src.rendering import render_report


def _prediction(*, long_text: str = "") -> InspectionPrediction:
    return InspectionPrediction(
        sample_id="synthetic-sample",
        summary=BridgeSummary(
            bridge_name="测试对象",
            bridge_id="ID-1",
            report_date="2026-08-03",
            overall_score="88.0",
            overall_grade="B级",
            overall_conclusion="总体状况良好",
        ),
        detailed_conclusion=(long_text or "结论文本",),
        recommendations=(
            Recommendation(index="1", category="尽快维修", content="修复构件", location="桥面"),
        ),
        defects=(
            DefectObservation(index="1", location="桥面", defect_type="裂缝", description="第一处"),
            DefectObservation(index="1", location="桥面", defect_type="裂缝", description="第二处"),
            DefectObservation(index="2", location="梁体", defect_type="露筋", description="第三处"),
        ),
        causes=(long_text or "雨水侵蚀",),
        treatments=(long_text or "及时修复",),
        safety_impact=(long_text or "影响通行安全",),
    )


def _all_text(document: Document) -> str:
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def test_renders_prediction_sections_and_reopens(tmp_path: Path) -> None:
    output = render_report(_prediction(), tmp_path / "prediction.docx")

    document = Document(output)
    text = _all_text(document)
    assert output.is_file()
    assert "概要表" in text
    assert "建议明细" in text
    assert "病害列表" in text
    assert "详细结论" in text
    assert "病害成因" in text
    assert "处置建议" in text
    assert "安全影响" in text
    assert "synthetic-sample" in text
    assert len(document.tables) == 3
    assert document.tables[1].rows[1].cells[0].text == "1"
    assert document.tables[2].rows[1].cells[3].text == "第一处"


def test_renders_gold_json_record_with_empty_lists(tmp_path: Path) -> None:
    record = {
        "sample_id": "empty-sample",
        "summary": {"bridge_name": "测试对象"},
        "detailed_conclusion": [],
        "recommendations": [],
        "defects": [],
        "causes": [],
        "treatments": [],
        "safety_impact": [],
    }
    output = render_report({"gold_version": 1, "records": [record]}, tmp_path / "empty.docx")

    document = Document(output)
    assert len(document.tables) == 3
    assert len(document.tables[1].rows) == 1
    assert len(document.tables[2].rows) == 1
    assert "empty-sample" in _all_text(document)


def test_merges_only_consecutive_equal_indexes(tmp_path: Path) -> None:
    output = render_report(_prediction(), tmp_path / "merged.docx")
    document = Document(output)
    table = document.tables[2]

    first = table.rows[1]._tr.tc_lst[0].tcPr.vMerge
    continuation = table.rows[2]._tr.tc_lst[0].tcPr.vMerge
    last = table.rows[3]._tr.tc_lst[0].tcPr.vMerge
    assert first is not None and first.val == "restart"
    assert continuation is not None and continuation.val == "continue"
    assert last is None
    assert table.rows[1].cells[0].text == "1"


def test_preserves_long_text_without_truncation(tmp_path: Path) -> None:
    long_text = "长文本-" + "关键事实。" * 1200
    output = render_report(_prediction(long_text=long_text), tmp_path / "long.docx")

    document = Document(output)
    assert long_text in _all_text(document)
