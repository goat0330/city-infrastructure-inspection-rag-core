from __future__ import annotations

import json
from pathlib import Path
import sys
import tempfile
import unittest

from docx import Document

REPO_ROOT = Path(__file__).parents[2]
sys.path.insert(0, str(REPO_ROOT))

from src.audit import audit_dataset  # noqa: E402


def write_label(path: Path, bridge_name: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    summary = document.add_table(rows=2, cols=2)
    summary.cell(0, 0).text = "项目"
    summary.cell(0, 1).text = "值"
    summary.cell(1, 0).text = "桥梁名称"
    summary.cell(1, 1).text = bridge_name
    recommendations = document.add_table(rows=2, cols=4)
    for cell, value in zip(recommendations.rows[0].cells, ("序号", "建议类别", "建议内容", "部位")):
        cell.text = value
    for cell, value in zip(recommendations.rows[1].cells, ("1", "尽快维修", "修复", "桥面")):
        cell.text = value
    defects = document.add_table(rows=2, cols=7)
    for cell, value in zip(defects.rows[0].cells, ("序号", "病害部位", "病害类型", "病害描述", "是否新增", "历史状态", "发展")):
        cell.text = value
    for cell, value in zip(defects.rows[1].cells, ("1", "桥面", "裂缝", "一条裂缝", "是", "无", "无")):
        cell.text = value
    document.add_paragraph("（1）详细结论")
    document.add_paragraph("结构状况良好")
    document.add_paragraph("（2）建议明细")
    document.add_paragraph("病害成因")
    document.add_paragraph("雨水侵蚀")
    document.add_paragraph("处置建议")
    document.add_paragraph("及时修复")
    document.add_paragraph("安全影响")
    document.add_paragraph("影响行车安全")
    document.save(path)


class DatasetAuditTests(unittest.TestCase):
    def test_pairing_and_legacy_doc_failure_are_explicit(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            labels = root / "labels"
            reports = root / "reports"
            for name in ("桥A", "桥B", "桥C", "桥D"):
                write_label(labels / f"2012年/{name}-无对比年度的信息提取报告.docx", name)
            (labels / "2012年/旧桥-无对比年度的信息提取报告.doc").parent.mkdir(parents=True, exist_ok=True)
            (labels / "2012年/旧桥-无对比年度的信息提取报告.doc").write_bytes(b"legacy")
            (reports / "桥A.docx").parent.mkdir(parents=True, exist_ok=True)
            (reports / "桥A.docx").write_bytes(b"report")
            (reports / "桥B-检测报告-修订.doc").write_bytes(b"report")
            (reports / "桥D-定检报告.docx").write_bytes(b"report")
            (reports / "桥D-报告.docx").write_bytes(b"report")

            report = audit_dataset(labels, reports)
            entries = {entry["label_relative_path"]: entry for entry in report["pairing"]["entries"]}
            self.assertEqual(entries["2012年/桥A-无对比年度的信息提取报告.docx"]["status"], "paired_exact")
            self.assertEqual(entries["2012年/桥B-无对比年度的信息提取报告.docx"]["status"], "paired_fuzzy")
            self.assertEqual(entries["2012年/桥D-无对比年度的信息提取报告.docx"]["status"], "duplicate")
            self.assertEqual(entries["2012年/桥C-无对比年度的信息提取报告.docx"]["status"], "missing")
            parse_entries = {entry["label_relative_path"]: entry for entry in report["label_parsing"]["entries"]}
            old_doc = parse_entries["2012年/旧桥-无对比年度的信息提取报告.doc"]
            self.assertEqual(old_doc["status"], "failed")
            self.assertEqual(old_doc["error_code"], "legacy_doc_unsupported")
            self.assertEqual(report["file_statistics"]["labels"]["by_extension"][".doc"], 1)
            self.assertEqual(report["label_parsing"]["failed"], 1)
            self.assertEqual(report["pairing"]["unmatched_report_count"], 0)

    def test_audit_cli_writes_json_without_absolute_input_paths(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            labels = root / "labels"
            reports = root / "reports"
            output = root / "output"
            write_label(labels / "桥A-无对比年度的信息提取报告.docx", "桥A")
            reports.mkdir()
            (reports / "桥A.docx").write_bytes(b"report")
            import subprocess

            result = subprocess.run(
                [
                    sys.executable,
                    str(REPO_ROOT / "scripts/audit_dataset.py"),
                    "--labels-dir",
                    str(labels),
                    "--reports-dir",
                    str(reports),
                    "--output-dir",
                    str(output),
                ],
                cwd=REPO_ROOT,
                check=True,
                capture_output=True,
                text=True,
            )
            saved = json.loads((output / "audit_report.json").read_text(encoding="utf-8"))
            self.assertEqual(saved["pairing"]["status_counts"]["paired_exact"], 1)
            self.assertNotIn(str(root), result.stdout)


if __name__ == "__main__":
    unittest.main()
