from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from src.submission import validate_submission


def _write_document(
    path: Path,
    *,
    headings: tuple[str, ...] = ("桥梁概要", "建议明细", "病害列表"),
    recommendation_rows: tuple[tuple[str, ...], ...] = (("1", "维修", "修复桥面", "桥面"),),
    defect_rows: tuple[tuple[str, ...], ...] = (("1", "桥面", "裂缝", "一条裂缝", "否", "无", "无"),),
) -> Path:
    document = Document()
    for heading in headings:
        document.add_heading(heading, level=1)

    summary = document.add_table(rows=3, cols=2)
    for cell, value in zip(summary.rows[0].cells, ("字段", "内容")):
        cell.text = value
    for row, values in zip(summary.rows[1:], (("桥梁名称", "示例桥"), ("总体评分", "90"))):
        for cell, value in zip(row.cells, values):
            cell.text = value

    recommendations = document.add_table(rows=1 + len(recommendation_rows), cols=4)
    for cell, value in zip(recommendations.rows[0].cells, ("序号", "建议类别", "建议内容", "病害部位")):
        cell.text = value
    for row, values in zip(recommendations.rows[1:], recommendation_rows):
        for cell, value in zip(row.cells, values):
            cell.text = value

    defects = document.add_table(rows=1 + len(defect_rows), cols=7)
    for cell, value in zip(
        defects.rows[0].cells,
        ("序号", "病害部位", "病害类型", "病害描述", "是否新增", "历史状态", "发展"),
    ):
        cell.text = value
    for row, values in zip(defects.rows[1:], defect_rows):
        for cell, value in zip(row.cells, values):
            cell.text = value

    document.save(path)
    return path


def _failure_codes(result: dict[str, object]) -> set[str]:
    return {str(item["code"]) for item in result["failures"]}  # type: ignore[index]


def test_valid_docx_is_deterministic_and_json_serializable(tmp_path: Path) -> None:
    path = _write_document(tmp_path / "submission.docx", defect_rows=(
        ("1", "桥面", "裂缝", "一条裂缝", "否", "无", "无"),
        ("2", "栏杆", "锈蚀", "局部锈蚀", "否", "无", "无"),
    ))

    first = validate_submission(path)
    second = validate_submission(path)

    assert first == second
    assert first["status"] == "passed"
    assert first["valid"] is True
    assert first["failures"] == []
    assert first["metrics"]["recommendation_rows"] == 1  # type: ignore[index]
    assert first["metrics"]["defect_rows"] == 2  # type: ignore[index]
    json.dumps(first, ensure_ascii=False, sort_keys=True)


def test_empty_bridge_id_is_allowed_for_valid_gold_summary(tmp_path: Path) -> None:
    path = _write_document(tmp_path / "gold-bridge-id-empty.docx")
    document = Document(str(path))
    bridge_id = document.tables[0].add_row().cells
    bridge_id[0].text = "桥梁编号"
    bridge_id[1].text = ""
    document.save(path)

    result = validate_submission(path)

    assert result["valid"] is True
    assert result["failures"] == []


def test_corrupt_docx_is_reported_without_silent_content_skip(tmp_path: Path) -> None:
    path = tmp_path / "broken.docx"
    path.write_bytes(b"not a docx package")

    result = validate_submission(path)

    assert result["valid"] is False
    assert "docx_read_failed" in _failure_codes(result)
    assert {check["status"] for check in result["checks"] if check["name"] == "headings"} == {"not_checked"}  # type: ignore[index]


def test_temporary_and_non_docx_names_fail_explicitly(tmp_path: Path) -> None:
    temporary = _write_document(tmp_path / "~$submission.docx")
    temporary_result = validate_submission(temporary)
    assert temporary_result["valid"] is False
    assert "temporary_file" in _failure_codes(temporary_result)

    other_extension = _write_document(tmp_path / "submission.doc")
    extension_result = validate_submission(other_extension)
    assert extension_result["valid"] is False
    assert "invalid_extension" in _failure_codes(extension_result)


def test_missing_heading_is_listed_even_when_tables_are_valid(tmp_path: Path) -> None:
    path = _write_document(tmp_path / "missing-heading.docx", headings=("桥梁概要", "建议明细"))

    result = validate_submission(path)

    assert result["valid"] is False
    assert any(
        item["code"] == "missing_heading" and item["details"]["section"] == "defects"  # type: ignore[index]
        for item in result["failures"]  # type: ignore[index]
    )


def test_empty_fields_are_reported_and_rows_are_counted(tmp_path: Path) -> None:
    path = _write_document(
        tmp_path / "empty-field.docx",
        recommendation_rows=(("1", "维修", "", "桥面"),),
        defect_rows=(("1", "", "裂缝", "一条裂缝", "否", "无", "无"),),
    )

    result = validate_submission(path)

    assert result["valid"] is False
    assert _failure_codes(result) >= {"empty_field"}
    assert result["metrics"]["recommendation_rows"] == 1  # type: ignore[index]
    assert result["metrics"]["defect_rows"] == 1  # type: ignore[index]
    fields = result["metrics"]["empty_fields"]  # type: ignore[index]
    assert {field["column"] for field in fields} >= {"content", "location"}


def test_empty_recommendation_and_defect_tables_fail_row_checks(tmp_path: Path) -> None:
    path = _write_document(tmp_path / "no-rows.docx", recommendation_rows=(), defect_rows=())

    result = validate_submission(path)

    assert result["valid"] is False
    assert _failure_codes(result) >= {"no_recommendation_rows", "no_defect_rows"}
    assert result["metrics"]["row_counts"] == {"recommendations": 0, "defects": 0}  # type: ignore[index]
